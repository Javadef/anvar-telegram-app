import logging
import os
from telegram import Update, InputFile
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from docx import Document
from docx.shared import Pt
from datetime import datetime, timedelta
from io import BytesIO

# Log yozish
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)

# Global o'zgaruvchilar
file_path = None
step = 0
newCarNum = ""
newName = ""
driver= ""
newRegion = ""
abym = ""

# Read the bot token from the environment variable
TOKEN = os.getenv("BOT_TOKEN")  # This will read the environment variable

if not TOKEN:
    raise ValueError("Bot token is not set in the environment variable")

  
months = {
        1: 'Январь', 2: 'Февраль', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'Октябрь', 11: 'Ноябрь', 12: 'Декабрь'
}


mo = {
    1: '01', 2: '02', 3: '03', 4: '04',
    5: '05', 6: '06', 7: '07', 8: '08',
    9: '09', 10: '10', 11: '11', 12: '12'
}

truck = {
    136: 'ShACMAN 01 136 QKA', 137: 'ShACMAN 01 137 QKA', 516: 'ShACMAN 01 516 PKA', 517: 'ShACMAN 01 517 PKA',
    429: 'ShACMAN 01 429 QKA', 430:  'ShACMAN 01 430 QKA', 431: 'ShACMAN 01 431 QKA', 67: 'MAN 01 067 OMA',
    324: 'MAN 01 324 GMA', 325: 'MAN 01 325 GMA', 643: 'ShACMAN 01 643 LKA',645: 'ShACMAN 01 645 LKA' ,
    725: 'ShACMAN 01 725 LKA' , 913: 'ShACMAN 01 913 OKA' , 914: 'ShACMAN 01 914 OKA' , 573: 'ShACMAN 01 573 LKA',
    574: 'ShACMAN 01 574 LKA'
    
    }


ddName = {
    136: 'Lapasov Shokir        ', 137: 'Tajiyev Mirfozil        ', 516: 'Axmedov Marat         ', 517: 'Jalolov Rustam       ',
    429: 'Umrzaqov Dilshod         ', 430:  'Kaimov Fayzulla   ', 431: 'Agzamov Iskandar    ', 67: 'Miryakubov Davron               ',
    324: '  Karimov Rixsiboy                             ', 325: 'Gulyamov Dilmurod    ', 643: 'Miraipov Baxtiyor     ',645: 'Akramov Farhod               ' ,
    725: 'Abdukarimov Temur   ' , 913: 'Karaxanov Baxromjon    ' , 914: 'Tulaganov Asatulla  ' , 573: 'Abdiraimov Shokir    ',
    574: 'Karimov Burxon           '
    }    
volume = {
      325: '88', 324: '88', 67: '88', 
      725: '47', 913: '47', 516: '47', 517: '47', 137 : '47', 430: '47',
      573: '56', 574: '56', 643: '56', 645: '56', 914 :'56', 136: '56', 429: '56', 431: '56'
     }



# Word faylini o'zgartirish funksiyasi
def change_text_in_docx(file_path: str, newCarNum: str, newName: str, newRegion: str, abym: str) -> BytesIO:
    doc = Document(file_path)

    today = datetime.today()
    matn_topildi = False  # Buni boshlang'ich qiymati False ga o'rnating


    mo = {i: f"{i:02}" for i in range(1, 13)}
    
    # Bugungi sanani olish
    day = today.day
    month = months[today.month]
    year = today.year
    formatted_date = f"         '{day}' {month} {year}"
    dA = (today + timedelta(days=1)).strftime('__%d__%m___%Y___')

    # O'zgartirishlar

    for table in doc.tables:
     for row in table.rows:
        for cell in row.cells:
            if 'REGION' in cell.text or 'carNum' in cell.text or 'date' in cell.text or 'dA' in cell.text or 'aby' in cell.text or 'Dname' in cell.text or 'kb' in cell.text:
                if 'REGION' in cell.text:
                    # Eski matnni o'zgartirish
                    cell.text = cell.text.replace('REGION', newRegion.upper())
                    # Shriftni belgilash
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)  # Region uchun 8 pt
                        run.font.name = 'Tahoma'

                if 'carNum' in cell.text:
                    # Mashina raqamini o'zgartirish
                    cell.text = cell.text.replace('carNum', newCarNum)
                    # Shriftni belgilash
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(10)  # carNum uchun 10 pt
                        run.font.name = 'Tahoma'
                    # Shriftni belgilash
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(10)  # carNum uchun 10 pt
                        run.font.name = 'Tahoma'        
                if 'Dname' in cell.text:  
                    # Mashina raqamini o'zgartirish
                    if newName == 'd':
                        newName = ddName[newCarNum]
                    cell.text = cell.text.replace('Dname', newName)
                    # Shriftni belgilash
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)  # carNum uchun 8 pt
                        run.font.name = 'Tahoma'  
                 

                if 'aby' in cell.text:
                    # Mashina raqamini o'zgartirish
                    cell.text = cell.text.replace('aby', abym)
                    # Shriftni belgilash
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)  # carNum uchun 8 pt
                        run.font.name = 'Tahoma'           

                if 'date' in cell.text:
                    # Bugungi sanani o'zgartirish
                    cell.text = cell.text.replace('date', formatted_date)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)  # Sana uchun 8 pt
                        run.font.name = 'Tahoma'
                        

                if 'dA' in cell.text:
                    # Ertangi sanani o'zgartirish
                    cell.text = cell.text.replace('dA', dA)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)  # Sana uchun 8 pt
                        run.font.name = 'Tahoma'
                if 'kb' in cell.text:
            # Eski matnni o'zgartirish
                    cell.text = cell.text.replace('kb', cub)
            # Shriftni belgilash
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)  # Region uchun 8 pt shrifti
                        run.font.name = 'Tahoma'            


                matn_topildi = True
    modified_file_stream = BytesIO()
    doc.save(modified_file_stream)
    modified_file_stream.seek(0)  # Kursorni boshiga qaytaramiz
    return modified_file_stream

# /start komandasi
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text('Salom! Shablon Word faylini yuboring.')

# Faylni qabul qilish
template_file_path = None

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    global template_file_path
    if update.message.document:
        file = await update.message.document.get_file()
        template_file_path = f'{update.message.document.file_id}.docx'
        await file.download_to_drive(template_file_path)
        await update.message.reply_text('Shablon qabul qilindi. Endi faylni o\'zgartirish uchun ma\'lumotlarni yuborishingiz mumkin.')


async def receive_variables(update: Update, context: ContextTypes.DEFAULT_TYPE,) -> None:
    
    user_data = context.user_data

    if 'step' not in user_data:
        user_data['step'] = 1

    if user_data['step'] == 1:
        newCarNum = update.message.text
        car = int(newCarNum)
        newCarNum = truck[car]
        driver = ddName[car]
        cuba = volume[car]
        global cub
        cub = cuba
    
        user_data.update({'newCarNum': newCarNum, 'driver': driver, 'kb': cub, 'step': 2})
        
        await update.message.reply_text('Shofyor ismini yuboring.')

    elif user_data['step'] == 2:
        newName = update.message.text
        user_data['newName'] = newName if newName != 'd' else user_data['driver']
        user_data['step'] = 3
        await update.message.reply_text('Region nomini yuboring.')

    elif user_data['step'] == 3:
        user_data['newRegion'] = update.message.text
        user_data['step'] = 4
        await update.message.reply_text('Abyomni yuboring.')

    elif user_data['step'] == 4:
        user_data['abym'] = update.message.text  # Kalitni tayinlash

        # O'zgartirish
        modified_file_stream = change_text_in_docx(template_file_path, user_data['newCarNum'], user_data['newName'], user_data['newRegion'], user_data['abym'])

        # O'zgartirilgan faylni yuborish
        await update.message.reply_document(InputFile(modified_file_stream, filename=f"{user_data['newCarNum']}_{user_data['newRegion']}.docx"))
        await update.message.reply_text('Yanami?')

        user_data['step'] = 5  # Qayta ishga tushirishga tayyor

    elif user_data['step'] == 5:
        if update.message.text.lower() == 'da':
            await update.message.reply_text("Qayta ishga tushirish...")
            user_data['step'] = 1  # Qayta boshlash uchun
        else:
            await update.message.reply_text("ok")

        # Global o'zgaruvchini tozalash
        user_data.clear()

# Botni ishga tushirish
def main():
    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, receive_variables))

    app.run_polling()

if __name__ == '__main__':
    main()
